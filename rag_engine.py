# rag_engine.py
import os
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"

import numpy as np
import docx
from io import BytesIO
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from langchain.embeddings import SentenceTransformerEmbeddings
from langchain.vectorstores import FAISS
from langchain.docstore.in_memory import InMemoryDocstore
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_together import ChatTogether
import faiss

# Load and preprocess the fixed file once
def load_fixed_document(filepath):
    document = docx.Document(filepath)
    full_text = "\n\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
    doc = LangchainDocument(page_content=full_text.strip())
    splitter = RecursiveCharacterTextSplitter(separators=["\n\n"], chunk_size=1500, chunk_overlap=500)
    return splitter.split_documents([doc])

# Build index once
def init_engine():
    chunks = load_fixed_document("data/FinancialData.docx")  # path to your fixed file
    embedding_model = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    texts = [doc.page_content for doc in chunks]
    embeddings = np.array(embedding_model.embed_documents(texts), dtype=np.float32)

    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)

    docstore = InMemoryDocstore({i: chunks[i] for i in range(len(chunks))})
    vector_store = FAISS(index, docstore, {}, embedding_model.embed_query)

    return index, docstore, embedding_model

index, docstore, embedding_model = init_engine()

def query_chatbot(question, chat_history):
    query_embedding = embedding_model.embed_query(question)
    D, I = index.search(np.array([query_embedding]), k=3)

    contexts = [docstore.search(idx).page_content for idx in I[0] if idx != -1]
    if not contexts:
        return "No relevant context found."

    context = "\n\n---\n\n".join(contexts)
    history_context = "\n".join([f"{msg['role'].capitalize()}: {msg['content']}" for msg in chat_history])

    chat_model = ChatTogether(
        together_api_key=os.getenv("TOGETHER_API_KEY"), 
        model="meta-llama/Llama-3.3-70B-Instruct-Turbo"
    )

    prompt = PromptTemplate(
    input_variables=["history", "context", "question"],
    template="""You are a financial advisor chatbot for Indian users. Use only the information provided in the knowledge base below to answer user queries. Do not make up information or use outside knowledge.

USER INPUT (if provided):
- Risk Profile: [low / medium / high]
- Investment Horizon: [e.g., 1 year / 5 years / 15 years]
- Amount to Invest: [e.g., ₹50,000 or ₹5,000/month]
- Liquidity Preference: [high / medium / low]
- Tax-Saving Needed: [yes / no]
- Special Conditions (if any): [senior citizen / saving for child / retirement planning]

TASK:
Recommend suitable investment options among Fixed Deposits, Savings Accounts, Public Provident Fund (PPF), Certificates of Deposit (CDs), and Mutual Funds.

For each recommended product, explain:
- Why it suits the user's input
- Expected returns
- Lock-in period
- Liquidity
- Tax implications
- Senior citizen benefits (if any)

Also mention when a product may *not* be ideal for the user's goals or constraints.

Rules:
1. If the user just greets, greet them back briefly (e.g., "Hello" → "Hi! I’m here to help with your financial investment queries").
2. If the question lacks detail, ask for specific information (risk profile, horizon, etc.).
3. Quote interest rates, returns, or other data *exactly* from the provided context.
4. You can make logical assumptions and give examples if something is not explicitly mentioned in the context.
5. If the user asks about a product in the context, explain it along with a detailed example using realistic figures. Do proper math.
6. Always state if the rate or benefit is specific to general or senior citizens.
7. If the user asks about mutual funds, answer only if context includes data.
9. If asked about maximum returns or minimum risks, choose from the different products in the context and give only one response.


Context: {context}

Chat History: {history}

Question: {question}

Strict rules:
- Use *only* the provided context.
- Do *not* include assumptions or outside knowledge.
- Keep replies concise and accurate.

Response:"""
)


    qa_chain = LLMChain(llm=chat_model, prompt=prompt)
    return qa_chain.run(history=history_context, context=context, question=question)
